import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report(output_path):
    doc = Document()
    
    # Title
    title = doc.add_heading('Project Report: Support Ticket Triage System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Author / Date
    author_para = doc.add_paragraph('Automated Issue Triage using Machine Learning')
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.style.font.size = Pt(12)
    
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "The Support Ticket Triage System is an automated, AI-powered web application "
        "designed to classify incoming support tickets and GitHub-style issues. "
        "The system instantly predicts the ticket's Category, Priority, and Intent using "
        "machine learning algorithms based on the ticket's title and description. "
        "This drastically reduces manual triage time and ensures critical issues are routed correctly."
    )
    
    doc.add_heading('2. System Architecture', level=1)
    doc.add_paragraph(
        "The system is built as a complete end-to-end pipeline, consisting of data processing, "
        "machine learning models, and a modern Streamlit web interface."
    )
    
    doc.add_heading('2.1. Tech Stack', level=2)
    doc.add_paragraph("• Frontend UI: Streamlit (with custom premium CSS, glassmorphism design)")
    doc.add_paragraph("• Machine Learning: Scikit-learn (TF-IDF, Logistic Regression, OneVsRest Classifier)")
    doc.add_paragraph("• Data Processing: Pandas, NumPy")
    doc.add_paragraph("• Artifact Persistence: Joblib")
    
    doc.add_heading('2.2. Processing Pipeline', level=2)
    doc.add_paragraph(
        "1. Input Handling: Accepts single text input (title + body) or batch CSV uploads.\n"
        "2. Preprocessing: Cleans text, normalizes whitespace, and combines title and body.\n"
        "3. Vectorization: Converts text to numerical features using TF-IDF (Term Frequency-Inverse Document Frequency) up to 50,000 features.\n"
        "4. Inference: Feeds vectorized text into three separate models to predict Category, Priority, and Intent."
    )
    
    doc.add_heading('3. Machine Learning Models', level=1)
    doc.add_paragraph(
        "The system employs robust baseline models optimized for text classification. "
        "All models have been trained locally and evaluated before deployment."
    )
    
    doc.add_heading('3.1. Category Model', level=2)
    doc.add_paragraph(
        "• Model Type: TF-IDF + OneVsRest Logistic Regression (Multi-label classification)\n"
        "• Dataset: Trained on 106,909 real GitHub issues.\n"
        "• Performance: Achieves a Micro F1 Score of 0.67 and Macro F1 of 0.60.\n"
        "• Output: Identifies up to 3 relevant categories (e.g., bug, feature_request, build_ci_cd) along with confidence scores."
    )
    
    doc.add_heading('3.2. Priority and Intent Models', level=2)
    doc.add_paragraph(
        "• Model Type: TF-IDF + Logistic Regression (Single-label classification)\n"
        "• Dataset: Trained on 8,469 customer support tickets.\n"
        "• Priority Classes: Critical, High, Medium, Low.\n"
        "• Intent Classes: Technical Issue, Billing Inquiry, Cancellation Request, Product Inquiry, Refund Request."
    )
    
    doc.add_heading('4. Web Interface & User Experience', level=1)
    doc.add_paragraph(
        "The user interface is developed using Streamlit, featuring a completely customized "
        "premium dark theme with modern design aesthetics."
    )
    doc.add_paragraph(
        "• Glassmorphism Design: Translucent, blurred card backgrounds for a sleek look.\n"
        "• Interactive Elements: Animated confidence bars and styled category chips.\n"
        "• Dual Input Modes: Users can either paste text for a quick single-ticket prediction "
        "or upload a CSV file to batch-predict hundreds of tickets at once.\n"
        "• Sidebar Analytics: Real-time metadata display showing the training rows, model types, "
        "and evaluation metrics for transparency."
    )
    
    doc.add_heading('5. Testing & Quality Assurance', level=1)
    doc.add_paragraph(
        "The project is fully covered by automated tests to ensure reliability:\n"
        "• Unit Tests: Verifying text preprocessing, missing value handling, and dataframe validation.\n"
        "• Integration Tests: Ensuring the full training pipeline successfully produces and loads model artifacts.\n"
        "• Code Quality: Managed by Ruff linter to enforce clean, standardized Python code."
    )
    
    doc.add_heading('6. Conclusion & Future Work', level=1)
    doc.add_paragraph(
        "The Support Ticket Triage System provides a reliable, fast, and visually stunning "
        "baseline for automated issue routing. The current Scikit-learn based implementation "
        "offers excellent performance and interpretability. Future enhancements may include "
        "fine-tuning Transformer models (e.g., DistilBERT) for improved accuracy and "
        "integrating directly with the GitHub API for real-time automated triage."
    )
    
    doc.save(output_path)
    print(f"Report successfully saved to: {output_path}")

if __name__ == "__main__":
    desktop_path = r"C:\Users\arsal\Desktop\Support_Ticket_Triage_System_Report.docx"
    create_report(desktop_path)
